import os
import io
from typing import List, Optional, Union
from PyPDF2 import PdfReader
from docx import Document
from openpyxl import load_workbook
import logging

logger = logging.getLogger(__name__)


class DocumentService:
    """
    文書処理サービス
    PDF、Word、Excelファイルからテキストを抽出する
    """

    def __init__(self):
        # 仕様に合わせ、Word(.docx)/Excel(.xlsx) のみ受付
        self.supported_extensions = ['docx', 'xlsx']
        self.max_file_size = 10 * 1024 * 1024  # 10MB

    async def extract_text_from_file(self, file_content: bytes, filename: str) -> str:
        """
        ファイルからテキスト抽出
        
        Args:
            file_content: ファイルのバイト列
            filename: ファイル名

        Returns:
            str: 抽出したテキスト

        Raises:
            ValueError: サポートされていないファイル形式の場合
            Exception: OCR処理エラーの場合
        """
        try:
            file_ext = filename.lower().split('.')[-1]

            if file_ext == 'docx':
                return await self._extract_from_docx(file_content)
            elif file_ext == 'xlsx':
                return await self._extract_from_xlsx(file_content)
            else: 
                raise ValueError(f'サポートされていないファイル形式: {file_ext}')
        
        except Exception as e:
            logger.error(f"OCR処理エラー({filename}): {e}")
            raise Exception(f"ファイルのテキスト抽出に失敗しました: {str(e)}")
        
    # async def _extract_from_pdf(self, pdf_content: bytes) -> str:
    #    """
    #    PDFファイルからテキスト抽出
    #    Args:
    #        pdf_content: PDFファイルの内容
    #    Returns:
    #        str: 抽出したテキスト
    #    """
    #    try:
    #        pdf_file = io.BytesIO(pdf_content)
    #        pdf_reader = PdfReader(pdf_file)

    #        all_text =[]
    #        for i, page in enumerate(pdf_reader.pages):
    #            page_text = page.extract_text()
    #            if page_text.strip():
    #                all_text.append(f'[ページ {i+1}]\n{page_text.strip()}')

    #        if not all_text:
    #            return '※PDFからテキストを抽出できませんでした'
            
    #        logger.info(f'PDF処理完了: {len(pdf_reader.pages)}ページ, 抽出文字数: {len("".join(all_text))}')
    #        return '\n\n'.join(all_text)
    #    except Exception as e:
    #        raise Exception(f"PDFファイルのテキスト抽出に失敗しました: {str(e)}")
        
    async def _extract_from_docx(self, docx_content: bytes) -> str:
        """
        Wordファイルからテキストを抽出
        
        Args:
            docx_content: Wordファイルの内容
            
        Returns:
            str: 抽出されたテキスト
        """
        try:
            docx_file = io.BytesIO(docx_content)
            doc = Document(docx_file)

            all_text = []

            # 段落のテキストを抽出
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    all_text.append(paragraph.text.strip())
            
            # 表のテキストを抽出
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(' | '.join(row_text))
                
                if table_text:
                    all_text.append('[表]\n' + '\n'.join(table_text))
            
            if not all_text:
                return "※Wordファイルからテキストを抽出できませんでした。"
            
            result_text = '\n\n'.join(all_text)
            logger.info(f"Word処理完了: 抽出文字数: {len(result_text)}")
            return result_text
            
        except Exception as e:
            raise Exception(f"Wordファイルからのテキスト抽出に失敗: {str(e)}")
        
    async def _extract_from_xlsx(self, xlsx_content: bytes) -> str:
        """
        Excelファイルからテキストを抽出
        
        Args:
            xlsx_content: Excelファイルの内容
            
        Returns:
            str: 抽出されたテキスト
        """
        try:
            xlsx_file = io.BytesIO(xlsx_content)
            workbook = load_workbook(xlsx_file, data_only=True)
            
            all_text = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = []
                
                # セルの内容を読み取り
                for row in sheet.iter_rows():
                    row_text = []
                    for cell in row:
                        if cell.value is not None:
                            cell_text = str(cell.value).strip()
                            if cell_text:
                                row_text.append(cell_text)
                    
                    if row_text:
                        sheet_text.append(' | '.join(row_text))
                
                if sheet_text:
                    all_text.append(f'[シート: {sheet_name}]\n' + '\n'.join(sheet_text))
            
            if not all_text:
                return "※Excelファイルからテキストを抽出できませんでした。"
            
            result_text = '\n\n'.join(all_text)
            logger.info(f"Excel処理完了: {len(workbook.sheetnames)} シート, 抽出文字数: {len(result_text)}")
            return result_text
            
        except Exception as e:
            raise Exception(f"Excelファイルからのテキスト抽出に失敗: {str(e)}")
    
    def validate_file(self, filename: str, file_size: int) -> bool:
        """
        ファイルのバリデーション

        Args:
            filename: ファイル名
            file_size: ファイルサイズ（バイト）

        Returns:
            bool: バリデーション結果

        Raises: 
            ValueError: バリデーションエラーの場合
        """
        # ファイル形式チェック
        file_ext = filename.lower().split('.')[-1]
        
        if file_ext not in self.supported_extensions:
            raise ValueError(f"サポートされていないファイル形式です。対応形式: PDF, Word(.docx), Excel(.xlsx)")
        
        # ファイルサイズチェック（10MB制限）
        if file_size > self.max_file_size:
            raise ValueError(f"ファイルサイズが大きすぎます。上限: {self.max_file_size // (1024*1024)}MB")
        
        return True
    
    def validate_file_count(self, current_count: int, new_files_count: int = 1) -> bool:
        """
        ファイル数のバリデーション
        
        Args:
            current_count: 現在のファイル数
            new_files_count: 追加するファイル数
            
        Returns:
            bool: バリデーション結果
            
        Raises:
            ValueError: バリデーションエラーの場合
        """
        max_files = 3
        
        if current_count + new_files_count > max_files:
            raise ValueError(f"アップロードできるファイル数は最大{max_files}つまでです。")
        
        return True